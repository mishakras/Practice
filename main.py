import sys
import os

from PyQt6 import QtGui, uic
from PyQt6 import QtWidgets

from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor

import design

class Word(QtWidgets.QMainWindow, design.Ui_Word):
    def __init__(self):
        super(Word, self).__init__()
        self.setupUi(self)
        self.CreateDoc.clicked.connect(self.Create_Document)
        self.ExerciseType.currentTextChanged.connect(self.Change_Subtype)
        self.Createexec.clicked.connect(self.Create_exercize)
        self.Finishblock.clicked.connect(self.Finish_block)
        self.CreateQuestion.clicked.connect(self.Create_Question)
        self.CreateAnswer.clicked.connect(self.Create_Answer)
        self.FinishQuestion.clicked.connect(self.Finish_Question)
        self.FinishExercize.clicked.connect(self.Finish_Exercize)
        self.doc=Document()
        self.exercize_count=1
        self.question_count=0
        self.UnitNum=""
        self.BlockNum=""
        self.BlockNumTotal=""
        self.UIType=""
        self.UISubtype=""
        self.paragraph=""
        self.answertext=""
        self.counter=0
        self.line=1
        self.CreateDoc.setEnabled(True)
        self.CreateQuestion.setEnabled(False)
        self.CreateAnswer.setEnabled(False)
        self.FinishQuestion.setEnabled(False)
        self.Createexec.setEnabled(False)
        self.FinishExercize.setEnabled(False)
        self.Finishblock.setEnabled(False)
    # Создаёт шапку документа
    def Create_Document(self):
        self.exercize_count=1
        self.question_count=0
        self.counter=0
        self.line=1
        if(len(self.UnitNumber.text())<2):
            self.UnitNum="0"+self.UnitNumber.text()
        else:
            self.UnitNum=self.UnitNumber.text()
        if(len(self.BlockNumber.text())<2):
            self.BlockNum="0"+self.BlockNumber.text()
        else:
            self.BlockNum=self.BlockNumber.text()
        if(len(self.BlockNumberTotal.text())<2):
            self.BlockNumTotal="0"+self.BlockNumberTotal.text()
        else:
            self.BlockNumTotal=self.BlockNumberTotal.text()
        paragraph = self.doc.add_paragraph()
        paragraph.style.font.size = Pt(12)
        paragraph.style.font.name = 'Calibri'
        run = paragraph.add_run("%TaskGroupInfoStart%")
        run = self.doc.add_paragraph().add_run("LearningCourse:"+self.CourseNumber.text()+" - "+self.CourseName.text())
        run = self.doc.add_paragraph().add_run("Unit:Unit "+self.UnitNum+". "+self.UnitName.text())
        run = self.doc.add_paragraph().add_run("Section:Section 01")
        run = self.doc.add_paragraph().add_run("Block:Block "+self.BlockNumTotal+". "+self.BlockType.currentText()+" "+self.BlockNum)
        run = self.doc.add_paragraph().add_run("TaskGroup:Task Group 01")
        run = self.doc.add_paragraph().add_run("%TaskGroupInfoEnd%")
        self.doc.add_paragraph()
        self.doc.add_paragraph("Unit "+self.UnitNumber.text()+" / "+self.UnitName.text()+" / "+self.BlockType.currentText()
        +" "+self.BlockNum+" / "+self.BlockName.text())
        self.Createexec.setEnabled(True)
        self.CreateDoc.setEnabled(False)
    def Change_Subtype(self):
        while(self.ExerciseSubType.count()>0):
            self.ExerciseSubType.removeItem(0)
        if(self.ExerciseType.currentText()=='Tick it'):
            self.ExerciseSubType.addItem("Правильный ответ один")
            self.ExerciseSubType.addItem("Правильных ответов несколько")
        elif(self.ExerciseType.currentText()=='Drag and drop'):
            self.ExerciseSubType.addItem("Единый текст")
            self.ExerciseSubType.addItem("Нумерованные предложения")
            self.ExerciseSubType.addItem("Нумерованные предложения, лишние слова")
    def Create_exercize(self):
        self.CreateQuestion.setEnabled(False)
        self.CreateAnswer.setEnabled(False)
        self.FinishQuestion.setEnabled(False)
        self.FinishExercize.setEnabled(False)
        self.Createexec.setEnabled(False)
        self.question_count=0
        if (self.BlockType.currentText()=="Reading"):
            BlockType="R"
        elif (self.BlockType.currentText()=="Video"):
            BlockType="L"
        elif (self.BlockType.currentText()=="Grammar"):
            BlockType="G"
        elif (self.BlockType.currentText()=="Vocabulary"):
            BlockType="V"
        elif (self.BlockType.currentText()=="Writing"):
            BlockType="W"
        paragraph = self.doc.add_paragraph()
        run = self.doc.add_paragraph(style="Heading 2").add_run("%ExerciseContentCode_%"+BlockType+self.BlockNumberTotal.text()+"."+str(self.exercize_count)+
        "%_ExerciseContentCode%"+self.ExerciseName.text())
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb=RGBColor(0x00, 0x00, 0x00)
        self.exercize_count=self.exercize_count+1
        if(self.ExerciseType.currentText()=='Tick it'):
            self.UIType="Choice"
            self.UISubtype="Radioline"
            self.CreateQuestion.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Writing'):
            self.UIType="Writing"
            self.UISubtype="Plain"
            self.CreateQuestion.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Speaking'):
            self.UIType="Speaking"
            self.UISubtype="Plain"
            self.FinishExercize.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Essay'):
            self.UIType="Essay"
            self.UISubtype="Essay"
            self.FinishExercize.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Matching'):
            self.UIType="Matching"
            self.UISubtype="Word"
            self.FinishExercize.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Categories'):
            self.UIType="Moving"
            self.UISubtype="Table"
            self.FinishExercize.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Drag and drop'):
            self.UIType="Moving"
            self.UISubtype="Inline"
            self.counter=0
            self.CreateQuestion.setEnabled(True)
        elif(self.ExerciseType.currentText()=='Multiple choice'):
            self.UIType="Choice"
            self.UISubtype="Select"
            self.CreateQuestion.setEnabled(True)
        run = self.doc.add_paragraph().add_run("%ExerciseInfoStart%")
        run = self.doc.add_paragraph().add_run("UIType:"+self.UIType)
        run = self.doc.add_paragraph().add_run("UISubtype:"+self.UISubtype)
        run = self.doc.add_paragraph().add_run("%PropertiesStart%")
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run("is_additional:")
        run = paragraph.add_run(self.Isadditional.currentText())
        run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
        if(self.ExerciseType.currentText()=='Tick it'):
            if(self.ExerciseSubType.currentText()=='Правильных ответов несколько'):
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run("multiple_answers:")
                run = paragraph.add_run("true")
                run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
        if(self.ExerciseType.currentText()=='Categories'):
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run("dont_check_places:")
            run = paragraph.add_run("true")
            run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
        if(self.ExerciseType.currentText()=='Drag and drop'):
            if(self.ExerciseSubType.currentText()=='Нумерованные предложения, лишние слова'):
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run("show_possible_answers_for_each_question:")
                run = paragraph.add_run("true")
                run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
            elif(self.ExerciseSubType.currentText()=='Единый текст'):
                paragraph = self.doc.add_paragraph()
                run = paragraph.add_run("display_questions_as_plain_text:")
                run = paragraph.add_run("true")
                run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
            paragraph = self.doc.add_paragraph()
            run = paragraph.add_run("show_answer_places:")
            run = paragraph.add_run("true")
            run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
        run = self.doc.add_paragraph().add_run("%PropertiesEnd%")
        run = self.doc.add_paragraph().add_run("%ExerciseInfoEnd%")
        self.doc.add_paragraph()
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run("%Task_%"+self.ExerciseText.toPlainText()+"%_Task%")
        run.font.italic=True
        run.font.cs_italic=True
        paragraph = self.doc.add_paragraph()
        run = self.doc.add_paragraph().add_run("%Questions_%")
        if(self.ExerciseType.currentText()=='Matching'):
            run = self.doc.add_paragraph().add_run("%Table_%")
            run = self.doc.add_paragraph()
        elif(self.ExerciseType.currentText()=='Categories'):
            run = self.doc.add_paragraph().add_run("%Table_%")
            run = self.doc.add_paragraph()
        while(self.AnswerType.count()>0):
            self.AnswerType.removeItem(0)
        if(self.UIType=='Choice'):
            self.AnswerType.addItem("Wrong")
            self.AnswerType.addItem("Right")
    def Create_Question(self):
        self.CreateQuestion.setEnabled(False)
        self.CreateAnswer.setEnabled(True)
        tempstr=self.Question.toPlainText()
        for i in range(len(tempstr)):
            if(tempstr[i-1]=="'"):
                tempstr=tempstr[:i-1]+"`"+tempstr[i:]
        self.Question.setPlainText(tempstr)
        if(self.UIType=='Choice'):
            if(self.AnswerType.count()==1):
                self.AnswerType.addItem("Right")
        if(self.UIType=='Choice'):
            if(self.question_count!=0):
                self.doc.add_paragraph()
            else:
                self.question_count=self.question_count+1
            run = self.doc.add_paragraph().add_run("%Question_%")
            run = self.doc.add_paragraph().add_run(self.Question.toPlainText())
        elif(self.ExerciseType.currentText()=='Drag and drop'):
            if(self.ExerciseSubType.currentText()=='Нумерованные предложения, лишние слова'):
                if(self.counter==0):
                    if(self.question_count!=0):
                        self.doc.add_paragraph()
                    else:
                        self.question_count=self.question_count+1
                    run = self.doc.add_paragraph().add_run("%Question_%")
                    self.counter=1
                else:
                    self.counter=2
                    self.paragraph=self.doc.add_paragraph(self.Question.toPlainText())
            else:
                self.paragraph=self.doc.add_paragraph(self.Question.toPlainText())
        elif(self.ExerciseType.currentText()=='Writing'):
            if(self.question_count!=0):
                self.doc.add_paragraph()
            else:
                self.question_count=self.question_count+1
            run = self.doc.add_paragraph().add_run("%Question_%")
            self.paragraph=self.doc.add_paragraph(self.Question.toPlainText())
    def Create_Answer(self):  
        tempstr=self.Answer.text()
        for i in range(len(tempstr)):
            if(tempstr[i-1]=="'"):
                tempstr=tempstr[:i-1]+"`"+tempstr[i:]
        self.Answer.clear()
        self.Answer.insert(tempstr)
        if(self.UIType=='Choice'):
            self.FinishQuestion.setEnabled(True)
            if(self.AnswerType.currentText()=="Right"):
                run = self.doc.add_paragraph().add_run(" ")
                paragraph = self.doc.add_paragraph(self.Answer.text()+" ", style='List Bullet 2')
                run=paragraph.add_run("(+)")
                run.font.color.rgb=RGBColor(0xff, 0x00, 0x00)
                if(self.ExerciseType.currentText()!='Tick it'):
                    self.AnswerType.removeItem(1)
                elif(self.ExerciseSubType.currentText()=="Правильный ответ один"):
                    self.AnswerType.removeItem(1)
            else:
                paragraph = self.doc.add_paragraph(self.Answer.text(), style='List Bullet 2')
        elif(self.ExerciseType.currentText()=='Drag and drop'):
            if(self.ExerciseSubType.currentText()=='Единый текст' or self.ExerciseSubType.currentText()=='Нумерованные предложения'):
                run = self.paragraph.add_run(" ")
                run=self.paragraph.add_run("("+self.Answer.text()+")")
                run.font.color.rgb=RGBColor(0xff, 0x00, 0x00)
                self.CreateAnswer.setEnabled(False)
                self.CreateQuestion.setEnabled(True)
                self.FinishExercize.setEnabled(True)
            else:
                if(self.counter==1):
                    run=self.doc.add_paragraph().add_run(self.Answer.text())
                    run.font.bold = True
                    self.CreateQuestion.setEnabled(True)
                else:
                    run = self.paragraph.add_run(" ")
                    run=self.paragraph.add_run("("+self.Answer.text()+")")
                    run.font.color.rgb=RGBColor(0xff, 0x00, 0x00)
                    self.FinishQuestion.setEnabled(True)
                    self.CreateAnswer.setEnabled(False)
                    self.counter=0
        elif(self.ExerciseType.currentText()=='Writing'):
            run = self.paragraph.add_run(" ")
            run = self.paragraph.add_run("("+self.Answer.text()+")")
            run.font.color.rgb=RGBColor(0xff, 0x00, 0x00)
            self.CreateAnswer.setEnabled(False)
            self.CreateQuestion.setEnabled(True)
            self.FinishExercize.setEnabled(True)
    def Finish_Question(self):
        self.FinishQuestion.setEnabled(False)
        if(self.ExerciseType.currentText()=='Drag and drop'):
            if(self.ExerciseSubType.currentText()=='Нумерованные предложения, лишние слова'):
                self.CreateAnswer.setEnabled(True)
            else:
                 self.CreateQuestion.setEnabled(True)
        else:
            self.CreateQuestion.setEnabled(True)
        run = self.doc.add_paragraph().add_run("%_Question%")
        self.CreateQuestion.setEnabled(True)
        self.FinishExercize.setEnabled(True)
    def Finish_Exercize(self):
        self.CreateQuestion.setEnabled(False)
        self.CreateAnswer.setEnabled(False)
        self.FinishExercize.setEnabled(False)
        if(self.ExerciseType.currentText()=='Matching'):
            self.doc.add_paragraph()
            run = self.doc.add_paragraph().add_run("%_Table%")
        elif(self.ExerciseType.currentText()=='Categories'):
            self.doc.add_paragraph()
            run = self.doc.add_paragraph().add_run("%_Table%")
        run = self.doc.add_paragraph().add_run("%_Questions%")
        run = self.doc.add_paragraph().add_run("%TotalMark_%"+self.TotalMark.text()+"%_TotalMark%")
        run.font.color.rgb=RGBColor(0x00, 0x00, 0xff)
        run.font.size = Pt(14)
        question_count=0
        self.doc.save(self.CourseNumber.text()+"_"+self.UnitNum+"_01_"+self.BlockNumTotal+"_01"+".docx")
        self.Createexec.setEnabled(True)
        self.Finishblock.setEnabled(True)
        self.TotalMark.clear()
        self.ExerciseName.clear()
        self.ExerciseText.setPlainText("")
    def Finish_block(self):
        self.doc=Document()
        self.Createexec.setEnabled(False)
        self.CreateDoc.setEnabled(True)
        self.Finishblock.setEnabled(False)
def main():
    app = QtWidgets.QApplication(sys.argv)  # Новый экземпляр QApplication
    window = Word()  # Создаём объект класса ExampleApp
    window.show()  # Показываем окно
    app.exec()  # и запускаем приложение

if __name__ == "__main__":
    main()
